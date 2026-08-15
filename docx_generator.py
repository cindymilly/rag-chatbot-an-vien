"""
docx_generator.py
-----------------
Module khởi tạo file Microsoft Word (.docx) chuẩn Văn bản Hành chính Việt Nam
theo quy định tại Nghị định 30/2020/NĐ-CP của Chính phủ:
- Phông chữ chuẩn: Times New Roman
- Cỡ chữ: Tiêu đề (14-15pt, Đậm), Tiêu đề mục (13pt, Đậm), Nội dung (13pt, Thường)
- Căn lề chuẩn: Trên 2cm, Dưới 2cm, Trái 3cm, Phải 1.5cm
- Tiêu ngữ Quốc hiệu & Tên cơ quan Công an xã An Viễn
"""

import io
import re
from datetime import datetime
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Thiết lập padding lề cho cell bảng"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_name, m_val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m_name}')
        node.set(qn('w:w'), str(m_val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_legal_docx(title: str, content_markdown: str, filename_prefix: str = "Phieu_Huong_Dan") -> bytes:
    """
    Chuyển đổi văn bản hướng dẫn Markdown thành file Word (.docx) chuẩn Nghị định 30/2020/NĐ-CP
    """
    doc = Document()

    # 1. Cấu hình Lề trang (Top 2cm, Bottom 2cm, Left 3cm, Right 1.5cm)
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.79)     # ~2.0 cm
        s.bottom_margin = Inches(0.79)  # ~2.0 cm
        s.left_margin = Inches(1.18)    # ~3.0 cm
        s.right_margin = Inches(0.59)   # ~1.5 cm

    # 2. Cấu hình Style Mặc định (Times New Roman, 13pt)
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(13)
    normal_style.font.color.rgb = RGBColor(15, 23, 42)  # Dark slate blue

    # 3. Tạo Quốc hiệu & Tên Cơ quan (Bảng 2 cột ẩn viền)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Cột trái: CÔNG AN TỈNH ĐỒNG NAI / CÔNG AN XÃ AN VIỄN
    cell_left = table.cell(0, 0)
    cell_left.width = Inches(3.2)
    p_l1 = cell_left.paragraphs[0]
    p_l1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l1 = p_l1.add_run("CÔNG AN TỈNH ĐỒNG NAI\n")
    r_l1.font.name = "Times New Roman"
    r_l1.font.size = Pt(11)
    
    r_l2 = p_l1.add_run("CÔNG AN XÃ AN VIỄN\n")
    r_l2.font.name = "Times New Roman"
    r_l2.font.size = Pt(11)
    r_l2.font.bold = True

    r_l3 = p_l1.add_run("________")
    r_l3.font.name = "Times New Roman"
    r_l3.font.size = Pt(10)

    # Cột phải: CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM / Độc lập - Tự do - Hạnh phúc
    cell_right = table.cell(0, 1)
    cell_right.width = Inches(3.8)
    p_r1 = cell_right.paragraphs[0]
    p_r1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_r1 = p_r1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    r_r1.font.name = "Times New Roman"
    r_r1.font.size = Pt(11)
    r_r1.font.bold = True

    r_r2 = p_r1.add_run("Độc lập - Tự do - Hạnh phúc\n")
    r_r2.font.name = "Times New Roman"
    r_r2.font.size = Pt(12)
    r_r2.font.bold = True

    r_r3 = p_r1.add_run("____________________")
    r_r3.font.name = "Times New Roman"
    r_r3.font.size = Pt(10)

    doc.add_paragraph()  # Khoảng trống

    # 4. Tên Loại Văn bản & Tiêu đề
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_title.add_run("PHIẾU HƯỚNG DẪN THỦ TỤC HÀNH CHÍNH\n")
    r_t.font.name = "Times New Roman"
    r_t.font.size = Pt(15)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(11, 83, 148)  # Deep Navy Blue

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(f"Ngày khởi tạo: {datetime.now().strftime('%d/%m/%Y %H:%M')} | Địa điểm: Công an xã An Viễn")
    r_sub.font.name = "Times New Roman"
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()  # Khoảng trống

    # 5. Xử lý Nội dung Văn bản từ Markdown
    lines = content_markdown.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Loại bỏ các ký tự Markdown thừa như ###, **, *, #
        if stripped.startswith('#'):
            # Tiêu đề mục (Heading)
            clean_head = re.sub(r'^[#\s]+', '', stripped).replace('**', '')
            p_h = doc.add_paragraph()
            p_h.paragraph_format.space_before = Pt(10)
            p_h.paragraph_format.space_after = Pt(4)
            r_h = p_h.add_run(clean_head)
            r_h.font.name = "Times New Roman"
            r_h.font.size = Pt(13)
            r_h.font.bold = True
            r_h.font.color.rgb = RGBColor(11, 83, 148)
        elif stripped.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            # Mục lớn
            clean_sec = stripped.replace('**', '')
            p_s = doc.add_paragraph()
            p_s.paragraph_format.space_before = Pt(8)
            p_s.paragraph_format.space_after = Pt(3)
            r_s = p_s.add_run(clean_sec)
            r_s.font.name = "Times New Roman"
            r_s.font.size = Pt(13)
            r_s.font.bold = True
        elif stripped.startswith(('-', '*', '•')):
            # Gạch đầu dòng
            clean_bullet = re.sub(r'^[-*•\s]+', '', stripped).replace('**', '')
            p_b = doc.add_paragraph(style='List Bullet')
            p_b.paragraph_format.space_after = Pt(2)
            p_b.paragraph_format.line_spacing = 1.15
            r_b = p_b.add_run(clean_bullet)
            r_b.font.name = "Times New Roman"
            r_b.font.size = Pt(13)
        else:
            # Đoạn văn bản thường
            clean_text = stripped.replace('**', '')
            p_t = doc.add_paragraph()
            p_t.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_t.paragraph_format.space_after = Pt(3)
            p_t.paragraph_format.line_spacing = 1.15
            r_t = p_t.add_run(clean_text)
            r_t.font.name = "Times New Roman"
            r_t.font.size = Pt(13)

    doc.add_paragraph()  # Khoảng trống

    # 6. Phần Chữ ký & Xác nhận (Chuẩn Văn bản Hành chính)
    table_sig = doc.add_table(rows=1, cols=2)
    table_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_sig.autofit = False

    cell_sig_l = table_sig.cell(0, 0)
    cell_sig_l.width = Inches(3.5)
    p_sl = cell_sig_l.paragraphs[0]
    r_sl1 = p_sl.add_run("Nơi nhận:\n")
    r_sl1.font.name = "Times New Roman"
    r_sl1.font.size = Pt(10)
    r_sl1.font.bold = True
    r_sl1.font.italic = True
    r_sl2 = p_sl.add_run("- Người yêu cầu;\n- Lưu: HS, HĐC.")
    r_sl2.font.name = "Times New Roman"
    r_sl2.font.size = Pt(10)
    r_sl2.font.italic = True

    cell_sig_r = table_sig.cell(0, 1)
    cell_sig_r.width = Inches(3.5)
    p_sr = cell_sig_r.paragraphs[0]
    p_sr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sr1 = p_sr.add_run("CÔNG AN XÃ AN VIỄN\n")
    r_sr1.font.name = "Times New Roman"
    r_sr1.font.size = Pt(12)
    r_sr1.font.bold = True
    r_sr2 = p_sr.add_run("(Hệ thống duyệt và cấp tự động)\n\n\n")
    r_sr2.font.name = "Times New Roman"
    r_sr2.font.size = Pt(10)
    r_sr2.font.italic = True

    # 7. Xuất Stream Bộ nhớ Bytes
    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream.getvalue()
